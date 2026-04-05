# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from openai import OpenAI
import math
from io import BytesIO

st.title("엑셀 - 워드 자동 변환기")

api_key = st.text_input("OpenAI API Key", type="password")
uploaded_file = st.file_uploader("Upload (Excel .xlsx)", type=["xlsx"])
title_input = st.text_input("Title", value="2024 족보")
file_name_input = st.text_input("File Name", value="제목 없음")


# ---------------- GPT 응답 파싱 ----------------
def get_text_from_response(response):
    try:
        return response.output[0].content[0].text
    except:
        try:
            return response.output_text
        except:
            return ""


# ---------------- 가로선 ----------------
def add_horizontal_line(paragraph):
    p = paragraph._element
    p_pr = p.get_or_add_pPr()

    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '000000')

    border.append(bottom)
    p_pr.append(border)


# ---------------- GPT 처리 ----------------
def process_question(number, q_text, client, doc):
    prompt = f"""
문제와 선지를 분리해주세요.
반드시 아래 형식으로 출력:
문제: ...
① ...
② ...
③ ...
④ ...
⑤ ...

입력:
{str(q_text)}
"""

    try:
        response = client.responses.create(
            model="gpt-5.3-chat-latest",
            input=prompt
        )
        content = get_text_from_response(response).strip()
    except Exception as e:
        st.error(e)
        content = "문제: 오류"

    lines = content.splitlines()
    question_line = ""
    choices = []

    for l in lines:
        s = l.strip()
        if s.startswith("문제:"):
            question_line = s.replace("문제:", "").strip()
        elif s.startswith(("①", "②", "③", "④", "⑤")):
            choices.append(s)

    if not question_line:
        question_line = lines[0] if lines else "복원 실패"

    try:
        para_q = doc.add_paragraph()
        para_q.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para_q.add_run(f"{number}. {question_line}").bold = True

        if choices:
            para_c = doc.add_paragraph()
            para_c.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for j, choice in enumerate(choices):
                if j > 0:
                    para_c.add_run().add_break(WD_BREAK.LINE)
                para_c.add_run(choice)

        doc.add_paragraph()

    except Exception as e:
        p = doc.add_paragraph()
        p.add_run(f"{number}. 복원 실패 ({e})").bold = True
        doc.add_paragraph()


# ---------------- 실행 ----------------
if st.button("Convert"):

    if not api_key:
        st.error("API Key 입력하세요")
        st.stop()

    if uploaded_file is None:
        st.error("엑셀 파일 업로드하세요")
        st.stop()

    with st.spinner("Generating..."):

        client = OpenAI(api_key=api_key)
        df = pd.read_excel(uploaded_file)

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(10)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

        # ---------------- 제목 ----------------
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title_para.add_run(title_input)
        run.bold = True
        run.underline = True
        run.font.size = Pt(16)
        doc.add_paragraph()

        # ---------------- 교수 표 ----------------
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        headers = ["교수님", "번호", "담당 강의 (시수)"]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.add_run(header).bold = True

        doc.add_paragraph()

        # ---------------- 문제 처리 ----------------
        max_q_number = 0

        # 🔥 사진자료 이후 제외
        if "사진 자료" in df.columns:
            photo_col_idx = df.columns.get_loc("사진 자료")
        else:
            photo_col_idx = df.shape[1]

        for col in range(2, photo_col_idx, 3):
            for i in range(len(df)):

                try:
                    number_raw = df.iloc[i, col]
                    q_text = df.iloc[i, col + 1]

                    try:
                        number = int(number_raw)
                    except:
                        number = str(number_raw) if not pd.isna(number_raw) else "?"

                    if isinstance(number, int):
                        max_q_number = max(max_q_number, number)

                    if pd.isna(q_text):
                        p = doc.add_paragraph()
                        p.add_run(f"{number}. 복원 실패").bold = True
                        doc.add_paragraph()
                        continue

                    process_question(number, q_text, client, doc)

                except Exception as e:
                    p = doc.add_paragraph()
                    p.add_run(f"복원 실패: {e}").bold = True
                    doc.add_paragraph()

        # ---------------- 정답 및 해설 ----------------
        p = doc.add_paragraph()
        add_horizontal_line(p)
        doc.add_paragraph("정답 및 해설").runs[0].bold = True

        total_rows = math.ceil(max_q_number / 5 * 2)
        table = doc.add_table(rows=total_rows, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 테두리
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')

        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            borders.append(border)

        tbl_pr.append(borders)

        # 내용 + 회색
        for i in range(total_rows):
            for j in range(5):
                cell = table.rows[i].cells[j]
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if i % 2 == 0:
                    value = i // 2 * 5 + j + 1
                    if value <= max_q_number:
                        para.add_run(str(value))

                    # 회색 배경
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:val"), "clear")
                    shading.set(qn("w:fill"), "D9D9D9")
                    cell._tc.get_or_add_tcPr().append(shading)

        # ---------------- 총평 ----------------
        for text in ["강의 총평", "족관 총평", "시험 난이도"]:
            p = doc.add_paragraph()
            add_horizontal_line(p)
            doc.add_paragraph(text).runs[0].bold = True

        # ---------------- 다운로드 ----------------
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Download",
            data=buffer.getvalue(),
            file_name=f"{file_name_input}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )